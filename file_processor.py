# file_processor.py
import pdfplumber
import docx
import pandas as pd
import os
from PIL import Image
import pytesseract
from transformers import BlipProcessor, BlipForConditionalGeneration
import cv2
import tempfile
from haystack import Document

class FileProcessor:
    @staticmethod
    def process_pdf(file_path):
        with pdfplumber.open(file_path) as pdf:
            text = "".join(page.extract_text() for page in pdf.pages if page.extract_text())
        return Document(content=text, meta={"source": "pdf"})

    @staticmethod
    def process_docx(file_path):
        doc = docx.Document(file_path)
        text = " ".join([para.text for para in doc.paragraphs if para.text])
        return Document(content=text, meta={"source": "docx"})

    @staticmethod
    def process_csv(file_path):
        df = pd.read_csv(file_path)
        text = " ".join(df.astype(str).agg(" ".join, axis=1))
        return Document(content=text, meta={"source": "csv"})

    @staticmethod
    def ocr_image(image_path):
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            return text.strip() if text.strip() else "No text detected in image."
        except Exception as e:
            return f"OCR failed: {str(e)}"

    @staticmethod
    def caption_image(image_path):
        try:
            processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base", use_fast=True)
            model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
            image = Image.open(image_path).convert("RGB")
            inputs = processor(images=image, return_tensors="pt")
            outputs = model.generate(**inputs)
            caption = processor.decode(outputs[0], skip_special_tokens=True)
            return caption
        except Exception as e:
            return f"Captioning failed: {str(e)}"

    @staticmethod
    def extract_frames(video_path, interval=5):
        try:
            vid = cv2.VideoCapture(video_path)
            if not vid.isOpened():
                return []
            frames = []
            count = 0
            fps = int(vid.get(cv2.CAP_PROP_FPS)) or 30
            while vid.isOpened():
                ret, frame = vid.read()
                if not ret:
                    break
                if count % (interval * fps) == 0: # Extract a frame every interval seconds.
                    temp_file = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
                    cv2.imwrite(temp_file.name, frame)
                    frames.append(temp_file.name)
                count += 1
            vid.release()
            return frames
        except Exception as e:
            print(f"Frame extraction failed: {str(e)}")
            return []

    @staticmethod
    def process_media(file_path):
        if file_path.lower().endswith((".jpg", ".png", ".jpeg")):
            ocr_text = FileProcessor.ocr_image(file_path)
            caption = FileProcessor.caption_image(file_path)
            text = f"OCR: {ocr_text}\nCaption: {caption}"
            return Document(content=text, meta={"source": "image"})
        
        elif file_path.lower().endswith((".mp4", ".avi", ".mov")):
            frames = FileProcessor.extract_frames(file_path)
            if not frames:
                return Document(content="No frames extracted from video.", meta={"source": "video"})
            texts = []
            for i, frame_path in enumerate(frames):
                ocr_text = FileProcessor.ocr_image(frame_path)
                caption = FileProcessor.caption_image(frame_path)
                texts.append(f"Frame {i+1} (at {i*5}s): OCR: {ocr_text}\nCaption: {caption}")
                os.remove(frame_path)
            text = "\n".join(texts)
            return Document(content=text, meta={"source": "video"})
        
        else:
            raise ValueError(f"Unsupported media type: {file_path}")

    def process_file(self, file_path):
        if file_path.endswith(".pdf"):
            return self.process_pdf(file_path)
        elif file_path.endswith(".docx"):
            return self.process_docx(file_path)
        elif file_path.endswith(".csv"):
            return self.process_csv(file_path)
        elif file_path.lower().endswith((".jpg", ".png", ".jpeg", ".mp4", ".avi", ".mov")):
            return self.process_media(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")