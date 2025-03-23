# main.py
import pdfplumber
import docx
import pandas as pd
import requests
import json
import os
from PIL import Image
import pytesseract
from transformers import BlipProcessor, BlipForConditionalGeneration
import cv2
import tempfile
from haystack import Document, Pipeline
from haystack_integrations.document_stores.chroma import ChromaDocumentStore
from haystack_integrations.components.retrievers.chroma import ChromaEmbeddingRetriever
from haystack.components.embedders import SentenceTransformersDocumentEmbedder, SentenceTransformersTextEmbedder
from haystack.components.writers import DocumentWriter

document_store = ChromaDocumentStore(persist_path="./chroma_data")
doc_embedder = SentenceTransformersDocumentEmbedder(model="sentence-transformers/all-MiniLM-L6-v2")
text_embedder = SentenceTransformersTextEmbedder(model="sentence-transformers/all-MiniLM-L6-v2")
retriever = ChromaEmbeddingRetriever(document_store=document_store)
doc_writer = DocumentWriter(document_store=document_store)
OLLAMA_URL = "http://localhost:11434/api/generate"

# --- Media Processing Functions ---
def ocr_image(image_path):
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text.strip() if text.strip() else "No text detected in image."
    except Exception as e:
        return f"OCR failed: {str(e)}"

def caption_image(image_path):
    try:
        processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
        model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
        image = Image.open(image_path).convert("RGB")
        inputs = processor(images=image, return_tensors="pt")
        outputs = model.generate(**inputs)
        caption = processor.decode(outputs[0], skip_special_tokens=True)
        return caption
    except Exception as e:
        return f"Captioning failed: {str(e)}"

def extract_frames(video_path, interval=5):
    try:
        vid = cv2.VideoCapture(video_path)
        if not vid.isOpened():
            return []
        frames = []
        count = 0
        fps = int(vid.get(cv2.CAP_PROP_FPS)) or 30
        while vid.isOpened():
            ret, frame = vid.read()
            if not ret:
                break
            if count % (interval * fps) == 0:
                temp_file = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
                cv2.imwrite(temp_file.name, frame)
                frames.append(temp_file.name)
            count += 1
        vid.release()
        return frames
    except Exception as e:
        print(f"Frame extraction failed: {str(e)}")
        return []

def process_media(file_path):
    if file_path.lower().endswith((".jpg", ".png", ".jpeg")):
        ocr_text = ocr_image(file_path)
        caption = caption_image(file_path)
        text = f"OCR: {ocr_text}\nCaption: {caption}"
        return Document(content=text, meta={"source": "image"})
    elif file_path.lower().endswith((".mp4", ".avi", ".mov")):
        frames = extract_frames(file_path)
        if not frames:
            return Document(content="No frames extracted from video.", meta={"source": "video"})
        texts = []
        for i, frame_path in enumerate(frames):
            ocr_text = ocr_image(frame_path)
            caption = caption_image(frame_path)
            texts.append(f"Frame {i+1} (at {i*5}s): OCR: {ocr_text}\nCaption: {caption}")
            os.remove(frame_path)  # Clean up
        text = "\n".join(texts)
        return Document(content=text, meta={"source": "video"})
    else:
        raise ValueError(f"Unsupported media type: {file_path}")

# --- Existing File Processing ---
def process_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        text = "".join(page.extract_text() for page in pdf.pages if page.extract_text())
    return Document(content=text, meta={"source": "pdf"})

def process_docx(file_path):
    doc = docx.Document(file_path)
    text = " ".join([para.text for para in doc.paragraphs if para.text])
    return Document(content=text, meta={"source": "docx"})

def process_csv(file_path):
    df = pd.read_csv(file_path)
    text = " ".join(df.astype(str).agg(" ".join, axis=1))
    return Document(content=text, meta={"source": "csv"})

def process_file(file_path):
    if file_path.endswith(".pdf"): return process_pdf(file_path)
    elif file_path.endswith(".docx"): return process_docx(file_path)
    elif file_path.endswith(".csv"): return process_csv(file_path)
    elif file_path.lower().endswith((".jpg", ".png", ".jpeg")) or file_path.lower().endswith((".mp4", ".avi", ".mov")):
        return process_media(file_path)
    else: raise ValueError(f"Unsupported file type: {file_path}")

# --- Pipelines ---
indexing_pipeline = Pipeline()
indexing_pipeline.add_component("doc_embedder", doc_embedder)
indexing_pipeline.add_component("doc_writer", doc_writer)
indexing_pipeline.connect("doc_embedder.documents", "doc_writer.documents")

query_pipeline = Pipeline()
query_pipeline.add_component("text_embedder", text_embedder)
query_pipeline.add_component("retriever", retriever)
query_pipeline.connect("text_embedder.embedding", "retriever.query_embedding")

def generate_summary(context, sources, prompt="Briefly summarize this text:"):
    source_hint = f"From {', '.join(sources)}: " if sources else ""
    full_prompt = f"{source_hint}{prompt} {context}"
    response = requests.post(
        OLLAMA_URL,
        json={"model": "mistral:7b-instruct-q4_0", "prompt": full_prompt, "max_tokens": 50, "temperature": 0.5, "top_p": 0.9},
        stream=True
    )
    full_response = ""
    for line in response.iter_lines():
        if line:
            chunk = json.loads(line.decode('utf-8'))
            full_response += chunk.get("response", "")
            if chunk.get("done", False): break
    return full_response.strip()

def main():
    # Process and index files (uncomment first run)
    files = [
        "data/sample.pdf",
        "data/sample.docx",
        "data/sample.csv",
        "data/images/dog.jpeg",  # Add your image
        "data/videos/dog_flipped.mp4"    # Add your video
    ]
    documents = [process_file(file) for file in files]
    indexing_pipeline.run({"doc_embedder": {"documents": documents}})

    # Test queries
    queries = [
        ("What’s in the files?", "Summarize the content:", []),
        ("Tell me about my resume", "Summarize my resume:", ["docx"]),
        ("What are the student grades?", "Summarize student performance:", ["csv"]),
        ("What’s the humanitarian data about?", "Summarize humanitarian data:", ["pdf"]),
        ("What’s in the image?", "Summarize image content:", ["image"]),
        ("What’s in the video?", "Summarize video content:", ["video"])
    ]

    for query, prompt, source_filter in queries:
        results = query_pipeline.run({"text_embedder": {"text": query}, "retriever": {"top_k": 3}})
        docs = results["retriever"]["documents"]

        if not source_filter:
            all_docs = query_pipeline.run({"text_embedder": {"text": query}, "retriever": {"top_k": 10}})["retriever"]["documents"]
            docs = []
            for source in ["pdf", "docx", "csv", "image", "video"]:
                source_doc = next((d for d in all_docs if d.meta["source"] == source), None)
                if source_doc: docs.append(source_doc)
        else:
            docs = [doc for doc in docs if doc.meta["source"] in source_filter]
        
        context = " ".join([doc.content for doc in docs])
        context = context[:500] + "..." if len(context) > 500 else context
        sources = list(set(doc.meta["source"] for doc in docs))
        response = generate_summary(context, sources, prompt)

        print("\n\nQuery:", query)
        print("Prompt:", f"{', '.join(sources)}: {prompt}" if sources else prompt)
        print("\nContext (truncated):", context)
        print("\nResponse:", response)

if __name__ == "__main__":
    main()